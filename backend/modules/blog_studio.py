"""Blog and article generation with multi-format export."""

from __future__ import annotations

import json
from io import BytesIO

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from ulid import ULID

from core.errors import NotFoundError, ProcessingError
from core.logging_config import get_logger
from core.model_router import model_router
from models.db_models import Blog, BlogFormat, BlogType, Project
from modules.content_engine import _get_project_context
from services.ollama_client import OllamaClient

log = get_logger(__name__)

_ollama = OllamaClient()

_SYSTEM_PROMPTS = {
    "blog": (
        "You are a travel blogger writing engaging, SEO-friendly blog posts. "
        "Include headings, subheadings, and a natural flow. Write in first person "
        "with vivid descriptions and practical tips."
    ),
    "guide": (
        "You are writing a comprehensive travel guide. Include logistics, costs, "
        "recommended itinerary, tips, and practical information. Structure with "
        "clear sections and bullet points where appropriate."
    ),
    "review": (
        "You are writing an honest, detailed travel review. Cover pros and cons, "
        "value for money, who it's best for, and specific recommendations. "
        "Be balanced and helpful."
    ),
    "trip_report": (
        "You are writing a day-by-day trip report. Include timeline, activities, "
        "highlights, challenges, and lessons learned. Keep a conversational, "
        "authentic tone."
    ),
}


async def generate_blog(
    db: AsyncSession,
    project_id: str,
    blog_type: str,
    context: dict | None = None,
) -> Blog:
    project = await db.get(Project, project_id)
    if not project:
        raise NotFoundError(f"Project {project_id} not found")

    if blog_type not in _SYSTEM_PROMPTS:
        raise ProcessingError(
            f"Invalid blog type: {blog_type}. Valid: {list(_SYSTEM_PROMPTS.keys())}"
        )

    project_context = await _get_project_context(db, project_id)
    if context:
        project_context.update(context)

    user_context = project_context.get("text", "")

    prompt = (
        f"Write a {blog_type.replace('_', ' ')} about this travel project.\n\n"
        f"Project: {project_context['project_name']}\n"
        f"Description: {project_context['project_description']}\n\n"
        f"Transcript from footage:\n{project_context['transcript_text'][:4000]}\n\n"
        f"Scene breakdown:\n{project_context['scenes_summary'][:1000]}\n\n"
    )
    if user_context:
        prompt += f"Additional instructions from the user:\n{user_context}\n\n"
    prompt += f"Write the complete {blog_type.replace('_', ' ')} in Markdown format."

    try:
        model = await model_router.get_model("blog")
        response = await _ollama.generate(
            model=model,
            prompt=prompt,
            system=_SYSTEM_PROMPTS[blog_type],
        )
    except Exception as exc:
        raise ProcessingError(f"Blog generation failed: {exc}") from exc

    word_count = len(response.split())
    title_line = response.split("\n", 1)[0].lstrip("# ").strip()
    title = title_line if title_line else f"{project.name} - {blog_type.replace('_', ' ').title()}"

    blog = Blog(
        id=str(ULID()),
        project_id=project_id,
        title=title,
        body=response,
        blog_type=BlogType(blog_type),
        format=BlogFormat.md,
        word_count=word_count,
    )
    db.add(blog)
    await db.commit()
    await db.refresh(blog)

    log.info("blog_generated", blog_id=blog.id, blog_type=blog_type, word_count=word_count)
    return blog


async def list_blogs(db: AsyncSession, project_id: str) -> list[Blog]:
    result = await db.execute(
        select(Blog)
        .where(Blog.project_id == project_id)
        .order_by(Blog.created_at.desc())
    )
    return list(result.scalars().all())


async def export_blog(db: AsyncSession, blog_id: str, format: str) -> bytes:
    blog = await db.get(Blog, blog_id)
    if not blog:
        raise NotFoundError(f"Blog {blog_id} not found")

    if format == "md":
        return blog.body.encode("utf-8")

    elif format == "html":
        try:
            import markdown

            html = markdown.markdown(blog.body, extensions=["extra", "toc"])
            full_html = (
                f"<!DOCTYPE html><html><head><meta charset='utf-8'>"
                f"<title>{blog.title}</title></head>"
                f"<body>{html}</body></html>"
            )
            return full_html.encode("utf-8")
        except ImportError:
            html = f"<pre>{blog.body}</pre>"
            return html.encode("utf-8")

    elif format == "docx":
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(blog.title, 0)
            for paragraph in blog.body.split("\n\n"):
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                if paragraph.startswith("# "):
                    doc.add_heading(paragraph.lstrip("# "), level=1)
                elif paragraph.startswith("## "):
                    doc.add_heading(paragraph.lstrip("## "), level=2)
                elif paragraph.startswith("### "):
                    doc.add_heading(paragraph.lstrip("### "), level=3)
                else:
                    doc.add_paragraph(paragraph)

            buffer = BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
        except ImportError:
            raise ProcessingError(
                "python-docx is not installed. Install with: pip install python-docx"
            )

    else:
        raise ProcessingError(f"Unsupported export format: {format}. Valid: md, html, docx")
